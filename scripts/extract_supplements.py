#!/usr/bin/env python3
"""Extract audit-friendly text/tables from PDF, DOCX, and XLSX supplements."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


def extract_pdf(path: Path, output_root: Path) -> dict[str, object]:
    reader = PdfReader(path)
    out = output_root / path.parent.name / f"{path.name}.txt"
    out.parent.mkdir(parents=True, exist_ok=True)
    parts = []
    for page_no, page in enumerate(reader.pages, start=1):
        parts.append(f"\n===== PAGE {page_no} =====\n")
        parts.append(page.extract_text() or "")
    out.write_text("".join(parts), encoding="utf-8")
    return {"source": path.as_posix(), "type": "pdf", "pages_or_sheets": len(reader.pages), "output": out.as_posix()}


def extract_docx(path: Path, output_root: Path) -> dict[str, object]:
    document = Document(path)
    out = output_root / path.parent.name / f"{path.name}.txt"
    out.parent.mkdir(parents=True, exist_ok=True)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table_no, table in enumerate(document.tables, start=1):
        parts.append(f"\n===== TABLE {table_no} =====")
        for row in table.rows:
            parts.append("\t".join(cell.text.replace("\n", " | ") for cell in row.cells))
    out.write_text("\n".join(parts) + "\n", encoding="utf-8")
    return {
        "source": path.as_posix(),
        "type": "docx",
        "pages_or_sheets": "",
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "output": out.as_posix(),
    }


def extract_xlsx(path: Path, output_root: Path) -> dict[str, object]:
    workbook = load_workbook(path, read_only=True, data_only=False)
    out_dir = output_root / path.parent.name / path.name
    out_dir.mkdir(parents=True, exist_ok=True)
    sheets = []
    for worksheet in workbook.worksheets:
        safe_name = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in worksheet.title)
        out = out_dir / f"{safe_name}.csv"
        with out.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            for row in worksheet.iter_rows(values_only=True):
                writer.writerow(["" if value is None else value for value in row])
        sheets.append({"name": worksheet.title, "rows": worksheet.max_row, "columns": worksheet.max_column, "output": out.as_posix()})
    return {"source": path.as_posix(), "type": "xlsx", "pages_or_sheets": len(sheets), "sheets": sheets}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-root", required=True, type=Path)
    parser.add_argument("--output-root", required=True, type=Path)
    parser.add_argument("--index", required=True, type=Path)
    args = parser.parse_args()

    records: list[dict[str, object]] = []
    for path in sorted(args.input_root.rglob("*")):
        if not path.is_file() or path.name.endswith(".meta.json"):
            continue
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            records.append(extract_pdf(path, args.output_root))
        elif suffix == ".docx":
            records.append(extract_docx(path, args.output_root))
        elif suffix == ".xlsx":
            records.append(extract_xlsx(path, args.output_root))
    args.index.parent.mkdir(parents=True, exist_ok=True)
    args.index.write_text(json.dumps(records, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
